"""导出格式生成 - JSON / CSV / PDF / Word / JPG

供 /api/export/ 使用：根据 format 参数返回不同文件格式
"""
import csv
import io
import json
from datetime import datetime
from decimal import Decimal

# 模块中文名映射（用于文件内标题/表头）
MODULE_NAMES = {
    "profile": "个人档案",
    "clothes": "衣物",
    "outfits": "穿搭",
    "recipes": "菜谱",
    "meals": "餐食",
    "shopping": "购物清单",
    "expenses": "记账",
    "tasks": "家务",
    "inventory": "库存",
    "trips": "行程",
    "trip_events": "行程事件",
    "commute": "通勤",
    "packing": "打包清单",
}


def _clean(data):
    """递归转 Decimal/datetime 为可序列化值"""
    if isinstance(data, Decimal):
        return float(data)
    if isinstance(data, (datetime,)):
        return data.isoformat()
    if isinstance(data, dict):
        return {k: _clean(v) for k, v in data.items()}
    if isinstance(data, (list, tuple)):
        return [_clean(v) for v in data]
    return data


def _module_rows(data):
    """返回 [(模块名, 行列表)]，行是 dict；过滤掉 version/exported_at/username 元信息"""
    out = []
    for key, rows in data.items():
        if key in ("version", "exported_at", "username"):
            continue
        if isinstance(rows, list) and rows:
            out.append((key, MODULE_NAMES.get(key, key), rows))
    return out


def to_json_bytes(data):
    return json.dumps(_clean(data), ensure_ascii=False, indent=2).encode("utf-8")


def to_csv_bytes(data):
    buf = io.StringIO()
    writer = csv.writer(buf)
    for key, name, rows in _module_rows(data):
        writer.writerow([f"== {name} =="])
        if rows:
            headers = list(rows[0].keys())
            writer.writerow(headers)
            for r in rows:
                writer.writerow([str(r.get(h, "")) for h in headers])
        writer.writerow([])
    return buf.getvalue().encode("utf-8-sig")


def to_pdf_bytes(data):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCN", parent=styles["Title"], fontSize=18)
    h2_style = ParagraphStyle("H2CN", parent=styles["Heading2"], fontSize=13, spaceBefore=12, spaceAfter=6)
    body_style = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontSize=9.5)

    story = [Paragraph("LifeHub 数据导出", title_style),
             Paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style),
             Spacer(1, 8)]
    for key, name, rows in _module_rows(data):
        story.append(Paragraph(name, h2_style))
        if not rows:
            story.append(Paragraph("(空)", body_style))
            continue
        headers = list(rows[0].keys())
        table_data = [[str(h) for h in headers]]
        for r in rows:
            table_data.append([str(r.get(h, ""))[:40] for h in headers])
        t = Table(table_data, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#7B6CF6")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7.5),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DDD")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F6FE")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(t)
        story.append(Spacer(1, 6))
    doc.build(story)
    return buf.getvalue()


def to_docx_bytes(data):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm

    doc = Document()
    doc.add_heading("LifeHub 数据导出", 0)
    doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for key, name, rows in _module_rows(data):
        doc.add_heading(name, level=1)
        if not rows:
            doc.add_paragraph("(空)")
            continue
        headers = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = str(h)
        for r in rows:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = str(r.get(h, ""))[:60]
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_jpg_bytes(data):
    """生成一张概览信息图 JPG（Pillow 绘制）"""
    from PIL import Image, ImageDraw, ImageFont

    # 找中文字体
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",       # 微软雅黑
        "C:/Windows/Fonts/simhei.ttf",     # 黑体
        "C:/Windows/Fonts/simsun.ttc",     # 宋体
    ]
    font_path = next((p for p in font_paths if __import__("os").path.exists(p)), None)

    def font(sz):
        return ImageFont.truetype(font_path, sz) if font_path else ImageFont.load_default()

    # 统计每模块条数
    lines = [("LifeHub 数据概览", 28, "#7B6CF6"), (datetime.now().strftime("%Y-%m-%d %H:%M"), 13, "#999")]
    total = 0
    for key, name, rows in _module_rows(data):
        lines.append((f"{name}: {len(rows)} 条", 16, "#333"))
        total += len(rows)
    lines.append(("", 8, "#fff"))
    lines.append((f"共 {total} 条数据", 18, "#FF006E"))

    W, H = 900, 120 + len(lines) * 34
    img = Image.new("RGB", (W, H), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    # 顶部渐变条
    for x in range(W):
        r = int(123 + (74 - 123) * x / W)
        g = int(108 + (144 - 108) * x / W)
        b = int(246 + (217 - 246) * x / W)
        draw.line([(x, 0), (x, 90)], fill=(r, g, b))
    y = 28
    for text, size, color in lines:
        draw.text((40, y), text, font=font(size), fill=color)
        y += 34
    buf = io.BytesIO()
    img.save(buf, "JPEG", quality=90)
    return buf.getvalue()


FORMAT_MAP = {
    "json": ("application/json", "json", to_json_bytes),
    "csv": ("text/csv", "csv", to_csv_bytes),
    "pdf": ("application/pdf", "pdf", to_pdf_bytes),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx", to_docx_bytes),
    "jpg": ("image/jpeg", "jpg", to_jpg_bytes),
}
